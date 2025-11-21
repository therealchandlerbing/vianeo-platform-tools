"""
Test suite for integration testing of players_influencers_generator

Tests cover:
- CLI interface and argument parsing
- End-to-end workflow (JSON → validation → DOCX)
- File I/O operations (loading, saving, error handling)
- Error recovery and edge cases
- Exit codes and error messages
"""

import pytest
import sys
import json
import subprocess
import tempfile
from pathlib import Path
from docx import Document

# Add tools directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "tools"))

from players_influencers_generator import (
    load_json_data,
    generate_document,
    main
)


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def fixtures_dir():
    """Return path to test fixtures directory"""
    return Path(__file__).parent / "fixtures"


@pytest.fixture
def tool_script():
    """Return path to the main tool script"""
    return Path(__file__).parent.parent / "tools" / "players_influencers_generator.py"


@pytest.fixture
def temp_output_file(tmp_path):
    """Create a temporary output file path"""
    return tmp_path / "test_output.docx"


# ============================================================================
# TESTS: File Loading (load_json_data)
# ============================================================================

@pytest.mark.integration
def test_load_json_data_valid_file(fixtures_dir):
    """Should successfully load valid JSON file"""
    json_path = fixtures_dir / "valid_input.json"
    data = load_json_data(str(json_path))

    assert data is not None
    assert "project_name" in data
    assert "players" in data
    assert "influencers" in data


@pytest.mark.integration
def test_load_json_data_missing_file():
    """Should exit with error when file doesn't exist"""
    with pytest.raises(SystemExit) as exc_info:
        load_json_data("nonexistent_file.json")

    assert exc_info.value.code == 1


@pytest.mark.integration
def test_load_json_data_malformed_json(fixtures_dir):
    """Should exit with error when JSON is malformed"""
    json_path = fixtures_dir / "malformed.json"

    with pytest.raises(SystemExit) as exc_info:
        load_json_data(str(json_path))

    assert exc_info.value.code == 1


@pytest.mark.integration
def test_load_json_data_not_json(fixtures_dir):
    """Should exit with error when file is not valid JSON"""
    txt_path = fixtures_dir / "not_json.txt"

    with pytest.raises(SystemExit) as exc_info:
        load_json_data(str(txt_path))

    assert exc_info.value.code == 1


@pytest.mark.integration
def test_load_json_data_empty_json(fixtures_dir):
    """Should load empty JSON object (will fail validation later)"""
    json_path = fixtures_dir / "empty.json"
    data = load_json_data(str(json_path))

    assert data == {}


# ============================================================================
# TESTS: End-to-End Workflow
# ============================================================================

@pytest.mark.integration
def test_end_to_end_workflow_with_example_file(fixtures_dir, temp_output_file):
    """Complete workflow: Load example JSON → Validate → Generate DOCX"""
    # Use the example file from the tools directory
    example_json = Path(__file__).parent.parent / "tools" / "example_rayla_ai.json"

    # Load data
    data = load_json_data(str(example_json))
    assert data is not None

    # Generate document
    result = generate_document(data, str(temp_output_file))
    assert result is True

    # Verify output file exists and is valid
    assert temp_output_file.exists()
    assert temp_output_file.stat().st_size > 0

    # Verify DOCX structure
    doc = Document(str(temp_output_file))
    assert len(doc.tables) == 2
    assert "Rayla AI" in '\n'.join([p.text for p in doc.paragraphs])


@pytest.mark.integration
def test_end_to_end_workflow_with_valid_fixture(fixtures_dir, temp_output_file):
    """Complete workflow with valid test fixture"""
    json_path = fixtures_dir / "valid_input.json"

    # Load data
    data = load_json_data(str(json_path))

    # Generate document
    result = generate_document(data, str(temp_output_file))
    assert result is True

    # Verify output
    assert temp_output_file.exists()
    doc = Document(str(temp_output_file))
    assert len(doc.tables) == 2


@pytest.mark.integration
def test_end_to_end_workflow_validation_failure(fixtures_dir, temp_output_file):
    """Workflow should stop at validation when data is invalid"""
    json_path = fixtures_dir / "too_few_players.json"

    # Load data (should succeed)
    data = load_json_data(str(json_path))

    # Generate document (should fail validation)
    result = generate_document(data, str(temp_output_file))
    assert result is False

    # Output file should not be created
    assert not temp_output_file.exists()


@pytest.mark.integration
def test_end_to_end_workflow_with_missing_fields(fixtures_dir, temp_output_file):
    """Workflow should fail gracefully with missing required fields"""
    json_path = fixtures_dir / "missing_project_name.json"

    data = load_json_data(str(json_path))
    result = generate_document(data, str(temp_output_file))

    assert result is False
    assert not temp_output_file.exists()


# ============================================================================
# TESTS: CLI Interface (Subprocess)
# ============================================================================

@pytest.mark.integration
def test_cli_with_no_arguments(tool_script):
    """CLI should show usage when called with no arguments"""
    result = subprocess.run(
        [sys.executable, str(tool_script)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "Usage:" in result.stdout or "usage:" in result.stdout.lower()


@pytest.mark.integration
def test_cli_with_one_argument(tool_script, fixtures_dir):
    """CLI should show usage when called with only one argument"""
    json_path = fixtures_dir / "valid_input.json"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "Usage:" in result.stdout or "usage:" in result.stdout.lower()


@pytest.mark.integration
def test_cli_with_valid_arguments(tool_script, fixtures_dir, tmp_path):
    """CLI should successfully generate document with valid arguments"""
    json_path = fixtures_dir / "valid_input.json"
    output_path = tmp_path / "cli_output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 0
    assert output_path.exists()
    assert "✓" in result.stdout or "complete" in result.stdout.lower()


@pytest.mark.integration
def test_cli_with_example_rayla_ai(tool_script, tmp_path):
    """CLI should work with the example Rayla AI file"""
    example_json = Path(__file__).parent.parent / "tools" / "example_rayla_ai.json"
    output_path = tmp_path / "rayla_output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(example_json), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 0
    assert output_path.exists()

    # Verify content
    doc = Document(str(output_path))
    assert len(doc.tables) == 2
    assert len(doc.tables[0].rows) == 10  # 1 header + 9 players
    assert len(doc.tables[1].rows) == 10  # 1 header + 9 influencers


@pytest.mark.integration
def test_cli_with_missing_input_file(tool_script, tmp_path):
    """CLI should exit with error when input file doesn't exist"""
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), "nonexistent.json", str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "not found" in result.stdout.lower() or "error" in result.stdout.lower()


@pytest.mark.integration
def test_cli_with_invalid_json(tool_script, fixtures_dir, tmp_path):
    """CLI should exit with error when JSON is malformed"""
    json_path = fixtures_dir / "malformed.json"
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "json" in result.stdout.lower() or "error" in result.stdout.lower()


@pytest.mark.integration
def test_cli_with_validation_failure(tool_script, fixtures_dir, tmp_path):
    """CLI should exit with error when validation fails"""
    json_path = fixtures_dir / "too_few_players.json"
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1
    assert "Validation" in result.stdout or "error" in result.stdout.lower()
    assert not output_path.exists()


# ============================================================================
# TESTS: Error Recovery and Edge Cases
# ============================================================================

@pytest.mark.integration
def test_output_file_path_with_spaces(fixtures_dir, tmp_path):
    """Should handle output paths with spaces"""
    json_path = fixtures_dir / "valid_input.json"
    output_path = tmp_path / "file with spaces.docx"

    data = load_json_data(str(json_path))
    result = generate_document(data, str(output_path))

    assert result is True
    assert output_path.exists()


@pytest.mark.integration
def test_output_file_in_nested_directory(fixtures_dir, tmp_path):
    """Should create output file in nested directories"""
    json_path = fixtures_dir / "valid_input.json"
    nested_dir = tmp_path / "subdir1" / "subdir2"
    nested_dir.mkdir(parents=True, exist_ok=True)
    output_path = nested_dir / "output.docx"

    data = load_json_data(str(json_path))
    result = generate_document(data, str(output_path))

    assert result is True
    assert output_path.exists()


@pytest.mark.integration
def test_overwrite_existing_file(fixtures_dir, tmp_path):
    """Should overwrite existing output file"""
    json_path = fixtures_dir / "valid_input.json"
    output_path = tmp_path / "output.docx"

    # Create initial file
    data = load_json_data(str(json_path))
    result1 = generate_document(data, str(output_path))
    assert result1 is True

    original_size = output_path.stat().st_size

    # Overwrite with same data
    result2 = generate_document(data, str(output_path))
    assert result2 is True

    # File should exist and have similar size
    assert output_path.exists()
    new_size = output_path.stat().st_size
    # Sizes should be close (DOCX files have some timestamp metadata)
    assert abs(new_size - original_size) < 1000


@pytest.mark.integration
def test_unicode_in_file_paths(fixtures_dir, tmp_path):
    """Should handle unicode characters in file paths"""
    json_path = fixtures_dir / "valid_input.json"
    output_path = tmp_path / "测试_output_café.docx"

    data = load_json_data(str(json_path))
    result = generate_document(data, str(output_path))

    assert result is True
    assert output_path.exists()


@pytest.mark.integration
def test_relative_vs_absolute_paths(fixtures_dir, tmp_path):
    """Should work with both relative and absolute paths"""
    json_path = fixtures_dir / "valid_input.json"

    # Test with absolute path
    abs_output = tmp_path / "abs_output.docx"
    data = load_json_data(str(json_path.absolute()))
    result1 = generate_document(data, str(abs_output.absolute()))
    assert result1 is True
    assert abs_output.exists()

    # Test with relative path (if possible)
    rel_output = tmp_path / "rel_output.docx"
    result2 = generate_document(data, str(rel_output))
    assert result2 is True
    assert rel_output.exists()


@pytest.mark.integration
def test_concurrent_file_generation(fixtures_dir, tmp_path):
    """Should handle generating multiple files in sequence"""
    json_path = fixtures_dir / "valid_input.json"
    data = load_json_data(str(json_path))

    # Generate multiple files
    output_files = []
    for i in range(5):
        output_path = tmp_path / f"output_{i}.docx"
        result = generate_document(data, str(output_path))
        assert result is True
        output_files.append(output_path)

    # Verify all files exist
    for output_file in output_files:
        assert output_file.exists()
        assert output_file.stat().st_size > 0


@pytest.mark.integration
def test_empty_json_object_handling(fixtures_dir, tmp_path):
    """Should handle empty JSON object gracefully"""
    json_path = fixtures_dir / "empty.json"
    output_path = tmp_path / "output.docx"

    data = load_json_data(str(json_path))
    result = generate_document(data, str(output_path))

    # Should fail validation
    assert result is False
    assert not output_path.exists()


# ============================================================================
# TESTS: Exit Codes and Error Messages
# ============================================================================

@pytest.mark.integration
def test_cli_exit_code_success(tool_script, fixtures_dir, tmp_path):
    """CLI should exit with code 0 on success"""
    json_path = fixtures_dir / "valid_input.json"
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 0


@pytest.mark.integration
def test_cli_exit_code_validation_failure(tool_script, fixtures_dir, tmp_path):
    """CLI should exit with code 1 on validation failure"""
    json_path = fixtures_dir / "invalid_acceptability.json"
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1


@pytest.mark.integration
def test_cli_exit_code_file_not_found(tool_script, tmp_path):
    """CLI should exit with code 1 when input file not found"""
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), "missing.json", str(output_path)],
        capture_output=True,
        text=True
    )

    assert result.returncode == 1


@pytest.mark.integration
def test_cli_provides_helpful_error_messages(tool_script, fixtures_dir, tmp_path):
    """CLI should provide clear error messages"""
    json_path = fixtures_dir / "too_few_players.json"
    output_path = tmp_path / "output.docx"

    result = subprocess.run(
        [sys.executable, str(tool_script), str(json_path), str(output_path)],
        capture_output=True,
        text=True
    )

    # Should mention validation errors
    output = result.stdout.lower()
    assert "validation" in output or "error" in output
    assert "players" in output or "8-10" in output


# ============================================================================
# TESTS: Real-World Scenarios
# ============================================================================

@pytest.mark.integration
@pytest.mark.slow
def test_full_workflow_with_all_acceptability_states(tmp_path):
    """Test complete workflow with data covering all acceptability states"""
    # Create comprehensive test data
    test_data = {
        "project_name": "Comprehensive Test Project",
        "executive_context": "Testing all acceptability states with realistic data.",
        "players": [
            {"name": "Favorable Player 1", "acceptability": "Favorable", "note": "Strong alignment"},
            {"name": "Favorable Player 2", "acceptability": "Favorable", "note": "Key supporter"},
            {"name": "Neutral Player 1", "acceptability": "Neutral", "note": "Undecided stance"},
            {"name": "Neutral Player 2", "acceptability": "Neutral", "note": "Wait-and-see"},
            {"name": "Unfavorable Player 1", "acceptability": "Unfavorable", "note": "Strong opposition"},
            {"name": "Unfavorable Player 2", "acceptability": "Unfavorable", "note": "Competitive threat"},
            {"name": "Unknown Player 1", "acceptability": "Don't know", "note": "Insufficient data"},
            {"name": "Unknown Player 2", "acceptability": "Don't know", "note": "Not yet engaged"},
        ],
        "influencers": [
            {"name": "Favorable Influencer 1", "acceptability": "Favorable", "note": "Advocate"},
            {"name": "Favorable Influencer 2", "acceptability": "Favorable", "note": "Champion"},
            {"name": "Neutral Influencer 1", "acceptability": "Neutral", "note": "Observer"},
            {"name": "Neutral Influencer 2", "acceptability": "Neutral", "note": "Monitoring"},
            {"name": "Unfavorable Influencer 1", "acceptability": "Unfavorable", "note": "Critic"},
            {"name": "Unfavorable Influencer 2", "acceptability": "Unfavorable", "note": "Skeptic"},
            {"name": "Unknown Influencer 1", "acceptability": "Don't know", "note": "New to scene"},
            {"name": "Unknown Influencer 2", "acceptability": "Don't know", "note": "Under research"},
        ]
    }

    output_path = tmp_path / "comprehensive_test.docx"
    result = generate_document(test_data, str(output_path))

    assert result is True
    assert output_path.exists()

    # Verify document structure
    doc = Document(str(output_path))
    assert len(doc.tables) == 2
    assert len(doc.tables[0].rows) == 9  # 8 players + header
    assert len(doc.tables[1].rows) == 9  # 8 influencers + header
