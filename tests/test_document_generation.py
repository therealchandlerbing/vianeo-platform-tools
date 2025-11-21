"""
Test suite for players_influencers_generator document generation

Tests cover:
- Successful DOCX file generation
- Document structure (sections, paragraphs, tables)
- Table structure (rows, columns, headers)
- Color coding for acceptability ratings
- Font styling (sizes, colors, bold/italic)
- Page layout (margins, page breaks)
- Error handling (file I/O, invalid paths)
"""

import pytest
import sys
import json
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Add tools directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "tools"))

from players_influencers_generator import (
    generate_document,
    COLORS,
    ACCEPTABILITY_CONFIG
)


# ============================================================================
# FIXTURES
# ============================================================================

@pytest.fixture
def valid_data():
    """Valid input data for document generation"""
    return {
        "project_name": "Test Project",
        "executive_context": "This is a test executive context for document generation.",
        "players": [
            {"name": "Player 1", "acceptability": "Favorable", "note": "Favorable note"},
            {"name": "Player 2", "acceptability": "Neutral", "note": "Neutral note"},
            {"name": "Player 3", "acceptability": "Unfavorable", "note": "Unfavorable note"},
            {"name": "Player 4", "acceptability": "Don't know", "note": "Unknown note"},
            {"name": "Player 5", "acceptability": "Favorable", "note": "Another favorable"},
            {"name": "Player 6", "acceptability": "Neutral", "note": "Another neutral"},
            {"name": "Player 7", "acceptability": "Unfavorable", "note": "Another unfavorable"},
            {"name": "Player 8", "acceptability": "Favorable", "note": "Final player"},
        ],
        "influencers": [
            {"name": "Influencer 1", "acceptability": "Favorable", "note": "Favorable influence"},
            {"name": "Influencer 2", "acceptability": "Neutral", "note": "Neutral influence"},
            {"name": "Influencer 3", "acceptability": "Unfavorable", "note": "Unfavorable influence"},
            {"name": "Influencer 4", "acceptability": "Don't know", "note": "Unknown influence"},
            {"name": "Influencer 5", "acceptability": "Favorable", "note": "Another favorable"},
            {"name": "Influencer 6", "acceptability": "Neutral", "note": "Another neutral"},
            {"name": "Influencer 7", "acceptability": "Unfavorable", "note": "Another unfavorable"},
            {"name": "Influencer 8", "acceptability": "Favorable", "note": "Final influencer"},
        ]
    }


@pytest.fixture
def temp_output_file(tmp_path):
    """Create a temporary output file path"""
    return tmp_path / "test_output.docx"


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def get_cell_background_color(cell):
    """Extract background color from cell shading"""
    shading_elms = cell._element.xpath('.//w:shd')
    if shading_elms:
        fill = shading_elms[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        if fill and fill != 'auto':
            # Convert hex to RGB tuple
            return tuple(int(fill[i:i+2], 16) for i in (0, 2, 4))
    return None


def get_text_color(run):
    """Extract text color from run"""
    if run.font.color.rgb:
        # Extract RGB values from color
        rgb = run.font.color.rgb
        return (rgb[0], rgb[1], rgb[2])
    return None


def rgb_to_hex(rgb_tuple):
    """Convert RGB tuple to hex string"""
    return '%02x%02x%02x' % rgb_tuple


# ============================================================================
# TESTS: Successful Generation
# ============================================================================

@pytest.mark.generation
def test_generate_document_creates_file(valid_data, temp_output_file):
    """Should successfully create a DOCX file"""
    result = generate_document(valid_data, str(temp_output_file))

    assert result is True
    assert temp_output_file.exists()
    assert temp_output_file.stat().st_size > 0


@pytest.mark.generation
def test_generate_document_returns_true_on_success(valid_data, temp_output_file):
    """Should return True when document is generated successfully"""
    result = generate_document(valid_data, str(temp_output_file))
    assert result is True


@pytest.mark.generation
def test_generate_document_returns_false_on_invalid_data(temp_output_file):
    """Should return False when validation fails"""
    invalid_data = {"project_name": "Test"}  # Missing required fields
    result = generate_document(invalid_data, str(temp_output_file))

    assert result is False
    assert not temp_output_file.exists()


# ============================================================================
# TESTS: Document Structure
# ============================================================================

@pytest.mark.generation
def test_generated_document_has_correct_sections(valid_data, temp_output_file):
    """Generated document should have correct number of sections"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    # Should have at least 1 section (default)
    assert len(doc.sections) >= 1


@pytest.mark.generation
def test_generated_document_has_two_tables(valid_data, temp_output_file):
    """Generated document should have exactly 2 tables (Players and Influencers)"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    assert len(doc.tables) == 2


@pytest.mark.generation
def test_generated_document_has_page_break(valid_data, temp_output_file):
    """Document should have a page break between Players and Influencers"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    # Check for page break in paragraphs
    has_page_break = False
    for para in doc.paragraphs:
        if para._element.xpath('.//w:br[@w:type="page"]'):
            has_page_break = True
            break

    assert has_page_break, "Document should contain a page break"


@pytest.mark.generation
def test_players_table_has_correct_structure(valid_data, temp_output_file):
    """Players table should have correct number of rows and columns"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]

    # Should have 3 columns (Entity, Acceptability, Strategic Implication)
    assert len(players_table.columns) == 3

    # Should have 1 header row + 8 data rows = 9 rows total
    assert len(players_table.rows) == 9


@pytest.mark.generation
def test_influencers_table_has_correct_structure(valid_data, temp_output_file):
    """Influencers table should have correct number of rows and columns"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    influencers_table = doc.tables[1]

    # Should have 3 columns
    assert len(influencers_table.columns) == 3

    # Should have 1 header row + 8 data rows = 9 rows total
    assert len(influencers_table.rows) == 9


@pytest.mark.generation
def test_table_headers_are_correct(valid_data, temp_output_file):
    """Table headers should have correct text"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    # Check Players table headers
    players_table = doc.tables[0]
    header_row = players_table.rows[0]

    assert "Players" in header_row.cells[0].text
    assert "Acceptability" in header_row.cells[1].text
    assert "Strategic Implication" in header_row.cells[2].text


# ============================================================================
# TESTS: Content Accuracy
# ============================================================================

@pytest.mark.generation
def test_document_contains_project_name(valid_data, temp_output_file):
    """Document should contain the project name in title"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    # Check first paragraph for project name
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    assert valid_data["project_name"] in full_text


@pytest.mark.generation
def test_document_contains_executive_context(valid_data, temp_output_file):
    """Document should contain the executive context"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    full_text = '\n'.join([para.text for para in doc.paragraphs])
    assert valid_data["executive_context"] in full_text


@pytest.mark.generation
def test_players_data_is_correct(valid_data, temp_output_file):
    """Players table should contain correct data"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]

    # Check first player (row 1, since row 0 is header)
    first_player_row = players_table.rows[1]
    assert valid_data["players"][0]["name"] in first_player_row.cells[0].text
    assert valid_data["players"][0]["acceptability"] in first_player_row.cells[1].text
    assert valid_data["players"][0]["note"] in first_player_row.cells[2].text


@pytest.mark.generation
def test_influencers_data_is_correct(valid_data, temp_output_file):
    """Influencers table should contain correct data"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    influencers_table = doc.tables[1]

    # Check first influencer
    first_influencer_row = influencers_table.rows[1]
    assert valid_data["influencers"][0]["name"] in first_influencer_row.cells[0].text
    assert valid_data["influencers"][0]["acceptability"] in first_influencer_row.cells[1].text
    assert valid_data["influencers"][0]["note"] in first_influencer_row.cells[2].text


# ============================================================================
# TESTS: Color Coding
# ============================================================================

@pytest.mark.generation
def test_favorable_acceptability_has_green_background(valid_data, temp_output_file):
    """Favorable acceptability cells should have green background"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    # First player has "Favorable" acceptability
    favorable_cell = players_table.rows[1].cells[1]

    bg_color = get_cell_background_color(favorable_cell)
    expected_color = COLORS["favorable_bg"]

    assert bg_color == expected_color, f"Expected {expected_color}, got {bg_color}"


@pytest.mark.generation
def test_neutral_acceptability_has_yellow_background(valid_data, temp_output_file):
    """Neutral acceptability cells should have yellow/amber background"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    # Second player has "Neutral" acceptability
    neutral_cell = players_table.rows[2].cells[1]

    bg_color = get_cell_background_color(neutral_cell)
    expected_color = COLORS["neutral_bg"]

    assert bg_color == expected_color, f"Expected {expected_color}, got {bg_color}"


@pytest.mark.generation
def test_unfavorable_acceptability_has_red_background(valid_data, temp_output_file):
    """Unfavorable acceptability cells should have red background"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    # Third player has "Unfavorable" acceptability
    unfavorable_cell = players_table.rows[3].cells[1]

    bg_color = get_cell_background_color(unfavorable_cell)
    expected_color = COLORS["unfavorable_bg"]

    assert bg_color == expected_color, f"Expected {expected_color}, got {bg_color}"


@pytest.mark.generation
def test_dont_know_acceptability_has_gray_background(valid_data, temp_output_file):
    """Don't know acceptability cells should have gray background"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    # Fourth player has "Don't know" acceptability
    dont_know_cell = players_table.rows[4].cells[1]

    bg_color = get_cell_background_color(dont_know_cell)
    expected_color = COLORS["dont_know_bg"]

    assert bg_color == expected_color, f"Expected {expected_color}, got {bg_color}"


# ============================================================================
# TESTS: Font and Styling
# ============================================================================

@pytest.mark.generation
def test_table_header_has_correct_background(valid_data, temp_output_file):
    """Table headers should have dark blue-gray background"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    header_cell = players_table.rows[0].cells[0]

    bg_color = get_cell_background_color(header_cell)
    expected_color = COLORS["table_header_bg"]

    assert bg_color == expected_color, f"Expected {expected_color}, got {bg_color}"


@pytest.mark.generation
def test_stakeholder_names_are_bold(valid_data, temp_output_file):
    """Stakeholder names in first column should be bold"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    players_table = doc.tables[0]
    name_cell = players_table.rows[1].cells[0]

    # Check if any run in the cell is bold
    has_bold = any(run.font.bold for run in name_cell.paragraphs[0].runs if run.text.strip())
    assert has_bold, "Stakeholder names should be bold"


# ============================================================================
# TESTS: Page Layout
# ============================================================================

@pytest.mark.generation
def test_document_has_correct_page_size(valid_data, temp_output_file):
    """Document should be letter size (8.5 x 11 inches)"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    section = doc.sections[0]

    # Check page width (8.5 inches = 7772400 EMU)
    assert abs(section.page_width - Inches(8.5)) < Inches(0.1)

    # Check page height (11 inches = 10058400 EMU)
    assert abs(section.page_height - Inches(11)) < Inches(0.1)


@pytest.mark.generation
def test_document_has_correct_margins(valid_data, temp_output_file):
    """Document should have 0.75 inch margins"""
    generate_document(valid_data, str(temp_output_file))
    doc = Document(str(temp_output_file))

    section = doc.sections[0]

    expected_margin = Inches(0.75)

    assert abs(section.top_margin - expected_margin) < Inches(0.01)
    assert abs(section.bottom_margin - expected_margin) < Inches(0.01)
    assert abs(section.left_margin - expected_margin) < Inches(0.01)
    assert abs(section.right_margin - expected_margin) < Inches(0.01)


# ============================================================================
# TESTS: Error Handling
# ============================================================================

@pytest.mark.generation
def test_generate_document_handles_invalid_path():
    """Should handle invalid output paths gracefully"""
    valid_data = {
        "project_name": "Test",
        "executive_context": "Context",
        "players": [{"name": f"P{i}", "acceptability": "Favorable", "note": f"N{i}"} for i in range(8)],
        "influencers": [{"name": f"I{i}", "acceptability": "Neutral", "note": f"N{i}"} for i in range(8)]
    }

    # Try to write to a directory that doesn't exist (with no permission to create)
    invalid_path = "/root/nonexistent/directory/output.docx"
    result = generate_document(valid_data, invalid_path)

    # Should return False on error
    assert result is False


@pytest.mark.generation
def test_generate_document_with_10_players_and_influencers(valid_data, temp_output_file):
    """Should handle maximum number of players and influencers (10 each)"""
    # Add 2 more to each list
    valid_data["players"].extend([
        {"name": "Player 9", "acceptability": "Favorable", "note": "Note 9"},
        {"name": "Player 10", "acceptability": "Neutral", "note": "Note 10"}
    ])
    valid_data["influencers"].extend([
        {"name": "Influencer 9", "acceptability": "Unfavorable", "note": "Note 9"},
        {"name": "Influencer 10", "acceptability": "Don't know", "note": "Note 10"}
    ])

    result = generate_document(valid_data, str(temp_output_file))
    assert result is True

    doc = Document(str(temp_output_file))
    assert len(doc.tables[0].rows) == 11  # 1 header + 10 players
    assert len(doc.tables[1].rows) == 11  # 1 header + 10 influencers


@pytest.mark.generation
def test_generate_document_with_unicode_characters(temp_output_file):
    """Should handle unicode characters in all text fields"""
    unicode_data = {
        "project_name": "项目名称 (Project Name) 🚀",
        "executive_context": "Context with émojis 🎯 and special characters: Café, Ñoño, 日本語",
        "players": [
            {"name": f"Player {i} 中文", "acceptability": "Favorable", "note": f"Note with café ☕"}
            for i in range(8)
        ],
        "influencers": [
            {"name": f"Influencer {i} العربية", "acceptability": "Neutral", "note": f"Note with Ñ"}
            for i in range(8)
        ]
    }

    result = generate_document(unicode_data, str(temp_output_file))
    assert result is True

    # Verify unicode content is preserved
    doc = Document(str(temp_output_file))
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    assert "🚀" in full_text or "中文" in full_text  # Should contain unicode
