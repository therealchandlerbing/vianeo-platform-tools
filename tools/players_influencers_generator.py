#!/usr/bin/env python3
"""
Players/Influencers Framework Generator
Version 2.0 | November 2025
Vianeo Business Model Evaluation Framework

Generates professional 2-page DOCX documents mapping ecosystem stakeholders
with acceptability ratings following BCG-level presentation standards.

Usage:
    python players_influencers_generator.py <input_json> <output_docx>

    Or import and use programmatically:
    from players_influencers_generator import generate_document
    generate_document(data, "output.docx")

Requirements:
    pip install python-docx
"""

import json
import sys
from typing import Dict, List, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ============================================================================
# COLOR PALETTE (BCG-Level Professional)
# ============================================================================

COLORS = {
    # Text Colors
    "title": (26, 26, 26),          # #1a1a1a - Near black
    "subtitle": (90, 90, 90),       # #5a5a5a - Medium gray
    "header": (44, 62, 80),         # #2c3e50 - Dark blue-gray
    "body": (58, 58, 58),           # #3a3a3a - Dark gray
    "entity": (44, 62, 80),         # #2c3e50 - Dark blue-gray
    "note": (74, 74, 74),           # #4a4a4a - Medium-dark gray

    # Table Headers
    "table_header_bg": (52, 73, 94),    # #34495e - Blue-gray
    "table_header_text": (255, 255, 255),  # #ffffff - White

    # Acceptability Color Coding
    "favorable_bg": (213, 244, 230),    # #d5f4e6 - Soft green
    "favorable_text": (39, 174, 96),    # #27ae60 - Green
    "neutral_bg": (255, 244, 217),      # #fff4d9 - Soft yellow
    "neutral_text": (243, 156, 18),     # #f39c12 - Amber
    "unfavorable_bg": (250, 219, 216),  # #fadbd8 - Soft red
    "unfavorable_text": (192, 57, 43),  # #c0392b - Red
    "dont_know_bg": (245, 245, 245),    # #f5f5f5 - Light gray
    "dont_know_text": (127, 140, 141),  # #7f8c8d - Medium gray

    # Borders
    "border_major": (189, 195, 199),    # #bdc3c7
    "border_minor": (236, 240, 241),    # #ecf0f1
}

# Acceptability rating configurations
ACCEPTABILITY_CONFIG = {
    "Favorable": {
        "bg": COLORS["favorable_bg"],
        "text": COLORS["favorable_text"]
    },
    "Neutral": {
        "bg": COLORS["neutral_bg"],
        "text": COLORS["neutral_text"]
    },
    "Unfavorable": {
        "bg": COLORS["unfavorable_bg"],
        "text": COLORS["unfavorable_text"]
    },
    "Don't know": {
        "bg": COLORS["dont_know_bg"],
        "text": COLORS["dont_know_text"]
    }
}


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_background(cell, rgb: tuple):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02x%02x%02x' % rgb)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """
    Set cell margins in DXA units (1 inch = 1440 DXA).
    Default: 100 DXA top/bottom (~0.07"), 140 DXA left/right (~0.1")
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for margin_name, value in [('top', top), ('bottom', bottom),
                                ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)


def set_cell_border(cell, color_rgb: tuple, sz=4, space=0):
    """Set cell border with specified color."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), '%02x%02x%02x' % color_rgb)
        tcBorders.append(border)

    tcPr.append(tcBorders)


def add_formatted_text(paragraph, text: str, font_size: int,
                       color_rgb: tuple, bold=False, italic=False):
    """Add formatted text to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold
    run.font.italic = italic
    return run


# ============================================================================
# VALIDATION FUNCTIONS
# ============================================================================

def validate_data(data: Dict[str, Any]) -> List[str]:
    """
    Validate input data structure and constraints.
    Returns list of validation errors (empty if valid).
    """
    errors = []

    # Required top-level fields
    required_fields = ["project_name", "executive_context", "players", "influencers"]
    for field in required_fields:
        if field not in data:
            errors.append(f"Missing required field: {field}")

    if errors:
        return errors  # Return early if structure is invalid

    # Validate project name
    if not data["project_name"] or len(data["project_name"]) > 100:
        errors.append("project_name must be 1-100 characters")

    # Validate executive context
    if not data["executive_context"] or len(data["executive_context"]) > 500:
        errors.append("executive_context must be 1-500 characters")

    # Validate players (8-10 entries)
    players = data.get("players", [])
    if len(players) < 8 or len(players) > 10:
        errors.append(f"Players must have 8-10 entries (found {len(players)})")

    for i, player in enumerate(players):
        errors.extend(validate_stakeholder(player, f"Player {i+1}"))

    # Validate influencers (8-10 entries)
    influencers = data.get("influencers", [])
    if len(influencers) < 8 or len(influencers) > 10:
        errors.append(f"Influencers must have 8-10 entries (found {len(influencers)})")

    for i, influencer in enumerate(influencers):
        errors.extend(validate_stakeholder(influencer, f"Influencer {i+1}"))

    return errors


def validate_stakeholder(stakeholder: Dict[str, str], prefix: str) -> List[str]:
    """Validate individual stakeholder entry."""
    errors = []

    # Required fields
    if "name" not in stakeholder:
        errors.append(f"{prefix}: Missing 'name' field")
    elif len(stakeholder["name"]) > 60:
        errors.append(f"{prefix}: Name exceeds 60 characters ({len(stakeholder['name'])})")

    if "acceptability" not in stakeholder:
        errors.append(f"{prefix}: Missing 'acceptability' field")
    elif stakeholder["acceptability"] not in ACCEPTABILITY_CONFIG:
        valid = ", ".join(ACCEPTABILITY_CONFIG.keys())
        errors.append(f"{prefix}: Invalid acceptability '{stakeholder['acceptability']}'. Must be one of: {valid}")

    if "note" not in stakeholder:
        errors.append(f"{prefix}: Missing 'note' field")
    elif len(stakeholder["note"]) > 250:
        errors.append(f"{prefix}: Note exceeds 250 characters ({len(stakeholder['note'])})")

    return errors


# ============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

def create_title_section(doc: Document, project_name: str, executive_context: str):
    """Create document title, subtitle, and executive context."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_formatted_text(title, f"{project_name}: Ecosystem Analysis",
                      font_size=48, color_rgb=COLORS["title"], bold=True)
    title.paragraph_format.space_after = Pt(6)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_formatted_text(subtitle, "Players & Influencers | Vianeo Framework",
                      font_size=24, color_rgb=COLORS["subtitle"], italic=True)
    subtitle.paragraph_format.space_after = Pt(12)

    # Executive Context
    context = doc.add_paragraph()
    context.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_formatted_text(context, executive_context,
                      font_size=20, color_rgb=COLORS["body"])
    context.paragraph_format.space_after = Pt(18)


def create_stakeholder_table(doc: Document, title: str, stakeholders: List[Dict[str, str]]):
    """Create formatted stakeholder table (Players or Influencers)."""
    # Section Header
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_formatted_text(header, title, font_size=28,
                      color_rgb=COLORS["header"], bold=True)
    header.paragraph_format.space_after = Pt(12)
    header.paragraph_format.space_before = Pt(6)

    # Create table (3 columns: Entity, Acceptability, Strategic Implication)
    table = doc.add_table(rows=1 + len(stakeholders), cols=3)
    table.style = 'Table Grid'

    # Set column widths (DXA units: 1 inch = 1440 DXA)
    # Column 1: 2400 DXA (~3.33"), Column 2: 1800 DXA (~2.5"), Column 3: 5160 DXA (~7.17")
    widths = [2400, 1800, 5160]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width * 635  # Convert DXA to EMU (914400 EMU = 1 inch)

    # Header row
    header_cells = table.rows[0].cells
    headers = [title.replace("Ecosystem ", ""), "Acceptability", "Strategic Implication"]

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, COLORS["table_header_bg"])
        set_cell_margins(cell, top=160, bottom=160, left=160, right=160)

        # Clear default paragraph and add formatted text
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        add_formatted_text(p, header_text, font_size=22,
                         color_rgb=COLORS["table_header_text"], bold=True)

    # Data rows
    for idx, stakeholder in enumerate(stakeholders):
        row = table.rows[idx + 1]

        # Entity name
        name_cell = row.cells[0]
        set_cell_margins(name_cell)
        set_cell_border(name_cell, COLORS["border_minor"])
        name_cell.text = ""
        p = name_cell.paragraphs[0]
        add_formatted_text(p, stakeholder["name"], font_size=20,
                         color_rgb=COLORS["entity"], bold=True)

        # Acceptability (color-coded)
        accept_cell = row.cells[1]
        acceptability = stakeholder["acceptability"]
        config = ACCEPTABILITY_CONFIG[acceptability]

        set_cell_background(accept_cell, config["bg"])
        set_cell_margins(accept_cell)
        set_cell_border(accept_cell, COLORS["border_minor"])

        accept_cell.text = ""
        p = accept_cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_formatted_text(p, acceptability, font_size=20,
                         color_rgb=config["text"], bold=True)

        # Strategic Note
        note_cell = row.cells[2]
        set_cell_margins(note_cell)
        set_cell_border(note_cell, COLORS["border_minor"])
        note_cell.text = ""
        p = note_cell.paragraphs[0]
        add_formatted_text(p, stakeholder["note"], font_size=19,
                         color_rgb=COLORS["note"])

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)


def generate_document(data: Dict[str, Any], output_path: str) -> bool:
    """
    Generate Players/Influencers DOCX document.

    Args:
        data: Dictionary containing project_name, executive_context, players, influencers
        output_path: Path to save the generated DOCX file

    Returns:
        True if successful, False otherwise
    """
    # Validate input data
    errors = validate_data(data)
    if errors:
        print("Validation Errors:")
        for error in errors:
            print(f"  - {error}")
        return False

    # Create document
    doc = Document()

    # Set document properties
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Create content
    create_title_section(doc, data["project_name"], data["executive_context"])
    create_stakeholder_table(doc, "Ecosystem Players", data["players"])

    # Page break before influencers
    doc.add_page_break()

    create_stakeholder_table(doc, "Ecosystem Influencers", data["influencers"])

    # Save document
    try:
        doc.save(output_path)
        print(f"✓ Successfully generated: {output_path}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False


# ============================================================================
# CLI INTERFACE
# ============================================================================

def load_json_data(file_path: str) -> Dict[str, Any]:
    """Load and parse JSON input file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in {file_path}: {e}")
        sys.exit(1)


def print_usage():
    """Print usage information."""
    print("""
Players/Influencers Framework Generator
Version 2.0 | Vianeo Business Model Evaluation Framework

Usage:
    python players_influencers_generator.py <input.json> <output.docx>

Input JSON Structure:
{
    "project_name": "Your Project Name",
    "executive_context": "2-3 sentence explanation of scope and interpretation",
    "players": [
        {
            "name": "Entity Name (max 60 chars)",
            "acceptability": "Favorable|Neutral|Unfavorable|Don't know",
            "note": "Strategic implication (max 250 chars)"
        }
        // ... 7-9 more (8-10 total)
    ],
    "influencers": [
        {
            "name": "Entity Name (max 60 chars)",
            "acceptability": "Favorable|Neutral|Unfavorable|Don't know",
            "note": "Strategic implication (max 250 chars)"
        }
        // ... 7-9 more (8-10 total)
    ]
}

Requirements:
    - 8-10 players (exactly)
    - 8-10 influencers (exactly)
    - Names: max 60 characters
    - Notes: max 250 characters
    - Valid acceptability values: Favorable, Neutral, Unfavorable, Don't know

Examples:
    python players_influencers_generator.py rayla_ai.json rayla_ai_output.docx
    """)


def main():
    """Main CLI entry point."""
    if len(sys.argv) != 3:
        print_usage()
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    print(f"Loading data from: {input_file}")
    data = load_json_data(input_file)

    print(f"Generating document: {output_file}")
    success = generate_document(data, output_file)

    if success:
        print("\n✓ Document generation complete!")
        print(f"\nDocument contains:")
        print(f"  - {len(data['players'])} Ecosystem Players")
        print(f"  - {len(data['influencers'])} Ecosystem Influencers")
        sys.exit(0)
    else:
        print("\n✗ Document generation failed.")
        sys.exit(1)


if __name__ == "__main__":
    main()
